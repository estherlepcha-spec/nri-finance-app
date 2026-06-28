from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
doc.core_properties.title = 'NRI Finance App User Manual'
doc.core_properties.subject = 'User guide, onboarding to offboarding, troubleshooting, and FAQs'
doc.core_properties.author = 'NRI Finance App Team'

def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

def add_paragraph(text, bold=False, italic=False, style=None):
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    if bold or italic:
        for run in p.runs:
            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True
    return p

def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def add_numbered(items):
    for item in items:
        doc.add_paragraph(item, style='List Number')

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    font = run.font
    font.name = 'Courier New'
    font.size = Pt(10)
    p.paragraph_format.left_indent = Pt(12)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F0F0F0')
    p._element.get_or_add_pPr().append(shading_elm)
    return p

def add_table(rows, cols):
    return doc.add_table(rows=rows, cols=cols)

def shade_cell(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_pPr().append(shading_elm)

# ============= COVER PAGE & INTRO =============
add_heading('NRI Finance App', level=0)
add_heading('Complete User Manual', level=1)
add_paragraph('Onboarding • Core Features • Troubleshooting • FAQs • Offboarding')
add_paragraph(' ')
add_paragraph('Version 1.0 | Last Updated: June 2026')
add_paragraph(' ')
add_heading('Welcome to Your Financial Freedom', level=2)
add_paragraph('The NRI Finance App is designed specifically for Non-Resident Indians, expatriates, and globally-mobile professionals who need to manage finances across multiple countries and currencies. This manual guides you from sign-up through every feature, including troubleshooting and FAQ sections.')

# ============= TABLE OF CONTENTS =============
add_heading('Table of Contents', level=1)
toc_items = [
    '1. Getting Started (Onboarding)',
    '2. Account Setup & Preferences',
    '3. Core Features Overview',
    '4. Day-to-Day Usage Guide',
    '5. Advanced Features',
    '6. Multi-Currency Management',
    '7. Family Finance Features',
    '8. Data Management & Security',
    '9. Troubleshooting Guide',
    '10. Frequently Asked Questions (FAQ)',
    '11. Offboarding & Data Export',
    '12. Contact & Support',
]
add_bullets(toc_items)

add_paragraph(' ')

# ============= SECTION 1: GETTING STARTED =============
add_heading('1. Getting Started (Onboarding)', level=1)

add_heading('1.1 Sign-Up Process', level=2)
add_numbered([
    'Visit the NRI Finance App website.',
    'Click "Sign Up" or "Get Started".',
    'Enter your email address and create a secure password.',
    'Verify your email via the link sent to your inbox.',
    'Log in to your new account.',
])

add_heading('1.2 Initial Setup Wizard', level=2)
add_paragraph('Upon first login, you\'ll see the Setup Wizard. This guides you through:')
add_bullets([
    'Selecting your home country (where you pay taxes/have family)',
    'Selecting your working country (where you currently earn)',
    'Choosing your primary currency (typically INR for NRIs)',
    'Connecting your first account (optional but recommended)',
    'Setting your monthly income (approximate)',
])
add_paragraph('You can skip any step and configure later in Settings.')

add_heading('1.3 First Login Checklist', level=2)
add_bullets([
    '✓ Email verified',
    '✓ Home and working countries selected',
    '✓ Home currency set (usually INR)',
    '✓ At least one account created',
    '✓ Profile preferences configured',
])

add_paragraph(' ')

# ============= SECTION 2: ACCOUNT SETUP =============
add_heading('2. Account Setup & Preferences', level=1)

add_heading('2.1 Creating Accounts', level=2)
add_paragraph('An "Account" in the app represents a bank account, credit card, or investment portfolio.')
add_numbered([
    'Navigate to "Accounts" from the sidebar.',
    'Click "Add New Account".',
    'Enter account name (e.g., "HDFC Savings Account", "Chase Checking").',
    'Select account type: Savings, Checking, Credit Card, Investment, or Other.',
    'Choose currency (INR, USD, GBP, EUR, etc.).',
    'Enter opening balance (from your bank statement).',
    'Click "Save".',
])

add_heading('2.2 Account Settings', level=2)
add_bullets([
    'Bank name (optional, for context)',
    'Account number (masked for security)',
    'Currency (multi-currency support)',
    'Opening balance date',
    'Account status (Active, Frozen, Closed)',
])

add_heading('2.3 User Preferences', level=2)
add_paragraph('Go to Settings → Preferences:')
add_bullets([
    'Home currency: Primary currency for net worth calculations',
    'Exchange rate source: Auto-update rates (recommended)',
    'Theme: Light or Dark mode',
    'Notifications: Push alerts for bills, reminders',
    'Language: English, Hindi (coming soon)',
    'Two-factor authentication: Enable for extra security',
])

add_paragraph(' ')

# ============= SECTION 3: CORE FEATURES =============
add_heading('3. Core Features Overview', level=1)

add_heading('3.1 Dashboard', level=2)
add_paragraph('Your financial snapshot at a glance:')
add_bullets([
    'Net Worth: Total assets minus liabilities in your home currency',
    'Monthly Income: All income sources this month',
    'Monthly Expenses: All spends this month',
    'Monthly Savings: Income minus expenses',
    'Remittances This Month: Money sent to India',
    'Upcoming Bills: Next 7 days of bills due',
    'Account Balances: Summary by currency',
])

add_heading('3.2 Transactions', level=2)
add_paragraph('Log every expense, income, and transfer:')
add_numbered([
    'Click "Add Transaction" or select an account.',
    'Choose transaction type: Expense, Income, or Transfer.',
    'Enter amount and select category (Groceries, Rent, Salary, etc.).',
    'Pick the date and account.',
    'Add notes (e.g., "Weekly groceries from Whole Foods").',
    'Optional: Attach a receipt image.',
])

add_heading('Transaction Categories:', level=3)
add_bullets([
    'Income: Salary, Freelance, Investment Returns, Gifts',
    'Expenses: Groceries, Rent, Utilities, Transport, Entertainment, Medical, Education, Fees',
    'Transfers: Internal (between your accounts) or External (to family, loans)',
])

add_heading('3.3 Smart Import (Bank Statements)', level=2)
add_paragraph('Import transactions directly from your bank:')
add_numbered([
    'Click "Transactions" → "Import".',
    'Upload your bank statement (PDF, Excel, CSV, or photo of the statement).',
    'The app auto-extracts transactions.',
    'Review and categorize suggested transactions.',
    'Click "Confirm Import" to add to your ledger.',
])

add_paragraph('The app automatically:')
add_bullets([
    'Categorizes expenses based on merchant and history',
    'Detects duplicate transactions',
    'Identifies transfers and remittances',
    'Flags anomalies (unusually high transactions)',
])

add_heading('3.4 Remittances', level=2)
add_paragraph('Track money sent to India or received from abroad:')
add_bullets([
    'Remittance: Amount sent home (tracked separately for efficiency analysis)',
    'Exchange rate: Locked rate at time of transfer',
    'Transfer receipt: Upload proof of transfer (TT receipt, online screenshot)',
    'Efficiency: Amount received vs. amount sent (shows conversion loss)',
    'Auto-link: App matches remittances to bank transfers automatically',
])

add_heading('Example Remittance Workflow:', level=3)
add_numbered([
    'Send $1,000 from your US account to India.',
    'Log remittance: "$1,000 at 83.5 INR/USD".',
    'Upload transfer receipt.',
    'When family deposits funds in India account, app auto-links them.',
    'View efficiency: "Sent $1,000 → Received ₹83,200" (after fees).',
])

add_heading('3.5 Bills & Recurring Payments', level=2)
add_paragraph('Never miss a bill again:')
add_numbered([
    'Click "Bills" → "Add Bill".',
    'Enter bill name, amount, due date, and account.',
    'Mark as recurring (monthly, quarterly, yearly) if applicable.',
    'Receive reminders 3 days, 1 day, and on due date.',
])

add_heading('3.6 Investments', level=2)
add_paragraph('Track stocks, mutual funds, FDs, and gold:')
add_numbered([
    'Click "Investments" → "Add Investment".',
    'Type: Stock, Mutual Fund, FD, Gold, Crypto, or Other.',
    'Enter quantity and current value.',
    'Track growth, dividends, and capital gains.',
])

add_heading('3.7 Loans & EMIs', level=2)
add_paragraph('Manage personal loans, home loans, and EMIs:')
add_bullets([
    'Loan name and type (Personal, Home, Auto, etc.)',
    'Principal amount, interest rate, and tenure',
    'EMI amount and due dates',
    'Payoff calculator: How much longer until loan is cleared?',
    'Extra payment tracking: When you pay more than EMI',
])

add_heading('3.8 Goals & Budget', level=2)
add_paragraph('Plan your financial future:')
add_numbered([
    'Set a goal: "Save $10,000 for India trip in 6 months".',
    'Track progress: App calculates required monthly savings.',
    'Get alerts: Notified if you\'re off track.',
])

add_paragraph(' ')

# ============= SECTION 4: DAY-TO-DAY USAGE =============
add_heading('4. Day-to-Day Usage Guide', level=1)

add_heading('4.1 Daily Habits', level=2)
add_bullets([
    'Check Dashboard each morning to see net worth and income/expense summary',
    'Log transactions immediately (or end of day) for accuracy',
    'Review upcoming bills to ensure funds available',
    'Update investment values if tracking manually',
])

add_heading('4.2 Weekly Tasks', level=2)
add_bullets([
    'Review week\'s spending vs. budget',
    'Check if any bills are due',
    'Reconcile accounts (compare app balance with bank statement)',
])

add_heading('4.3 Monthly Ritual', level=2)
add_heading('First of the month:', level=3)
add_numbered([
    'Review previous month\'s summary on Dashboard',
    'Check if goals are on track',
    'Update investment valuations',
    'Plan remittance (if applicable)',
])

add_heading('Mid-month:', level=3)
add_numbered([
    'Import latest bank statement',
    'Categorize any unreviewed transactions',
    'Check upcoming bills for next 15 days',
])

add_heading('End of month:', level=3)
add_numbered([
    'Review total income, expenses, and savings',
    'Log remittance if sending money home',
    'Update budget if needed',
    'Export summary for tax purposes (if needed)',
])

add_paragraph(' ')

# ============= SECTION 5: ADVANCED FEATURES =============
add_heading('5. Advanced Features', level=1)

add_heading('5.1 Estelle (AI Financial Advisor)', level=2)
add_paragraph('Chat with our AI advisor powered by Claude:')
add_bullets([
    'Ask financial questions: "Should I invest in mutual funds?"',
    'Get insights: "Where am I spending the most?"',
    'Plan scenarios: "If I save $500/month, when can I buy a car?"',
    'Extract documents: Upload bank statements, investment documents for auto-analysis',
])

add_heading('5.2 Smart Rules & Auto-Categorization', level=2)
add_paragraph('Create custom rules to auto-categorize transactions:')
add_numbered([
    'Go to Settings → Smart Rules.',
    'Click "Add Rule".',
    'Define: If merchant contains "Amazon", categorize as "Shopping".',
    'Apply rule to existing and future transactions.',
])

add_heading('5.3 Balance Audit & Reconciliation', level=2)
add_paragraph('Verify your app balances match your bank statements:')
add_numbered([
    'Click on any account.',
    'Click "Audit Balance".',
    'Enter current bank balance.',
    'App shows discrepancies and helps find missing transactions.',
])

add_heading('5.4 Tax Estimator', level=2)
add_paragraph('Calculate worldwide tax liability across 13 countries:')
add_bullets([
    'Income by country (US salary, India rental income, etc.)',
    'Deductions (donations, business expenses, etc.)',
    'Get estimated tax to pay and potential refunds',
])

add_heading('5.5 What-If Simulator', level=2)
add_paragraph('Scenario planning for major financial decisions:')
add_bullets([
    'Simulate: "If I increase savings by 20%, net worth in 5 years?"',
    'Test: "Impact of salary cut or bonus?"',
    'Plan: "Move to new country: income/expense changes?"',
])

add_paragraph(' ')

# ============= SECTION 6: MULTI-CURRENCY MANAGEMENT =============
add_heading('6. Multi-Currency Management', level=1)

add_heading('6.1 How Multi-Currency Works', level=2)
add_paragraph('All accounts can be in different currencies. The app automatically converts to your home currency (INR) for net worth calculations.')

add_heading('Example:', level=3)
add_bullets([
    'Account 1: $5,000 USD (US savings)',
    'Account 2: ₹50,000 INR (India account)',
    'Net Worth = (5000 × 83.5) + 50000 = 467,500 INR',
])

add_heading('6.2 Exchange Rates', level=2)
add_paragraph('The app uses live exchange rates from market sources:')
add_bullets([
    'Auto-update: Rates refresh every hour',
    'Manual rate: Override with custom rates (for locked-in transfers)',
    'History: See historical rates for any date',
])

add_heading('6.3 Transfer Between Currencies', level=2)
add_numbered([
    'Go to any account → "Transfer".',
    'Select destination account (different currency).',
    'Enter amount in source currency.',
    'Enter exchange rate (or use current rate).',
    'App calculates amount received in destination currency.',
])

add_heading('Example Transfer:', level=3)
add_code('From: $1,000 USD\nExchange Rate: 1 USD = 83.5 INR\nTo: ₹83,500 INR (approximately, minus bank fees)')

add_paragraph(' ')

# ============= SECTION 7: FAMILY FINANCE FEATURES =============
add_heading('7. Family Finance Features', level=1)

add_heading('7.1 Add Family Members', level=2)
add_numbered([
    'Go to Settings → Family.',
    'Click "Add Family Member".',
    'Enter name, email, relationship (Spouse, Child, Parent, etc.).',
    'Set permissions: View-only or Full access.',
    'Send invitation email.',
])

add_heading('7.2 Shared Budgets', level=2)
add_paragraph('Create budgets visible to family members:')
add_bullets([
    'Household expenses budget (rent, groceries, utilities)',
    'Education budget for children',
    'Healthcare budget for family',
])

add_heading('7.3 View Family Net Worth', level=2)
add_paragraph('Aggregate view of entire family\'s finances:')
add_bullets([
    'Combined net worth across all family members',
    'Who\'s spending most? (helpful for household budgeting)',
    'Family savings rate and goals progress',
])

add_paragraph(' ')

# ============= SECTION 8: DATA MANAGEMENT & SECURITY =============
add_heading('8. Data Management & Security', level=1)

add_heading('8.1 Your Data Is Secure', level=2)
add_bullets([
    'End-to-end encryption: Data encrypted in transit and at rest',
    'Password security: Never stored in plain text',
    'Two-factor authentication: Optional but recommended',
    'Privacy: We never sell or share your data',
    'GDPR compliant: Your data, your control',
])

add_heading('8.2 Backup & Sync', level=2)
add_bullets([
    'Automatic backup: All data synced to secure cloud',
    'Real-time sync: Changes on your phone/laptop sync instantly',
    'Offline mode: Access app offline, syncs when back online',
])

add_heading('8.3 Change Password', level=2)
add_numbered([
    'Go to Settings → Security.',
    'Click "Change Password".',
    'Enter current password and new password (min 12 characters).',
    'Click "Update".',
])

add_heading('8.4 Enable Two-Factor Authentication (2FA)', level=2)
add_numbered([
    'Settings → Security → Enable 2FA.',
    'Choose: SMS or Authenticator app.',
    'Verify your method.',
    'Save backup codes in a safe place (lost device recovery).',
])

add_heading('8.5 Export Your Data', level=2)
add_paragraph('Download all your financial data as CSV or Excel:')
add_numbered([
    'Settings → Data Management → Export.',
    'Choose format: CSV, Excel, or PDF report.',
    'Select date range.',
    'Click Export. File downloads to your computer.',
])

add_paragraph(' ')

# ============= SECTION 9: TROUBLESHOOTING =============
add_heading('9. Troubleshooting Guide', level=1)

add_heading('9.1 Login Issues', level=2)

add_heading('Problem: "Incorrect email or password"', level=3)
add_bullets([
    'Double-check spelling of your email',
    'Password is case-sensitive',
    'Caps Lock may be on',
    'Try: Click "Forgot Password" to reset',
])

add_heading('Problem: "Email not verified"', level=3)
add_bullets([
    'Check your email inbox and spam folder',
    'Look for email from noreply@nrifinance.app',
    'Click the verification link in the email',
    'Try: Resend verification email (link on login page)',
])

add_heading('9.2 Account & Balance Issues', level=2)

add_heading('Problem: "My app balance doesn\'t match my bank"', level=3)
add_bullets([
    'Use the Balance Audit tool (see Section 5.3)',
    'Check for pending transactions (not yet cleared by bank)',
    'Verify opening balance was entered correctly',
    'Look for un-categorized or missing imports',
    'Try: Manually reconcile from the account page',
])

add_heading('Problem: "I added a transaction but don\'t see it"', level=3)
add_bullets([
    'Refresh the page (Ctrl+R or Cmd+R)',
    'Check if it went to the correct account and date',
    'Use search to find it (Ctrl+F)',
    'If still missing: Try logging out and back in',
])

add_heading('9.3 Import & Sync Issues', level=2)

add_heading('Problem: "Bank statement import failed"', level=3)
add_bullets([
    'Ensure file format is PDF, Excel, or CSV (not image)',
    'Check file size is under 10 MB',
    'Try again with a different file format from your bank',
    'If recurring issue: Contact support with file sample',
])

add_heading('Problem: "Changes not syncing to my other device"', level=3)
add_bullets([
    'Check internet connection on both devices',
    'Try signing out and back in on the other device',
    'Wait 30 seconds for sync to occur',
    'If still not synced: Restart the app',
])

add_heading('9.4 Exchange Rate & Currency Issues', level=2)

add_heading('Problem: "Exchange rates are outdated"', level=3)
add_bullets([
    'Check Settings → Preferences → Exchange Rate Source',
    'Ensure "Auto-update" is enabled',
    'Rates update every hour automatically',
    'Manual rates: You can override for specific transactions',
])

add_heading('Problem: "Wrong net worth due to incorrect rate"', level=3)
add_bullets([
    'Check Settings → Currency Conversion',
    'Verify which rate is being used for each conversion',
    'Update rate manually if needed (Settings → Update Rates)',
])

add_heading('9.5 Performance & App Issues', level=2)

add_heading('Problem: "App is slow or freezing"', level=3)
add_bullets([
    'Close other browser tabs to free up memory',
    'Clear browser cache (Settings in your browser)',
    'Try a different browser (Chrome, Safari, Firefox, Edge)',
    'Check internet speed (slow internet = slow app)',
])

add_heading('Problem: "I see a blank screen or errors"', level=3)
add_bullets([
    'Hard refresh: Ctrl+Shift+R (Windows) or Cmd+Shift+R (Mac)',
    'Clear cookies and cache, then refresh',
    'Try in Incognito/Private browsing mode',
    'Check browser console for error messages (F12)',
])

add_heading('9.6 Security & Access Issues', level=2)

add_heading('Problem: "I forgot my password"', level=3)
add_numbered([
    'Click "Forgot Password" on the login page.',
    'Enter your email.',
    'Check email for password reset link.',
    'Click link and set new password.',
    'Try logging in with new password.',
])

add_heading('Problem: "I lost my 2FA phone"', level=3)
add_bullets([
    'Use backup codes you saved earlier',
    'Contact support with proof of identity',
    'We\'ll help you disable 2FA and secure your account',
])

add_heading('9.7 Feature-Specific Issues', level=2)

add_heading('Problem: "Remittance not linking to deposit"', level=3)
add_bullets([
    'Check dates: Remittance send date vs. deposit date (may be 3-5 days)',
    'Verify amounts match (within exchange rate variance)',
    'Try manual linking: Click on deposit → Link to remittance',
    'Check if transfer was to different account/person',
])

add_heading('Problem: "Bills not reminding me"', level=3)
add_bullets([
    'Ensure notifications are enabled in Settings',
    'Check browser notification permissions',
    'Verify bill due date is correct',
    'Look at notification history to see past alerts',
])

add_paragraph(' ')

# ============= SECTION 10: FAQ =============
add_heading('10. Frequently Asked Questions (FAQ)', level=1)

add_heading('Getting Started', level=2)

qa_pairs = [
    ('Is the app free?', 'Yes, basic features are free. Premium features (family sharing, advanced tax planning) require a subscription.'),
    ('What countries are supported?', 'We support 50+ countries and all major currencies. Home countries optimized: India, USA, UK, Canada, Australia, UAE, Singapore.'),
    ('Can I use the app on my phone?', 'Yes, the app is fully responsive. Use it on any device: laptop, tablet, or smartphone. Progressive Web App (PWA) support for offline access coming soon.'),
    ('Is my data safe?', 'Yes. All data is encrypted, stored securely on Supabase (SOC 2 Type II certified), and we comply with GDPR and privacy laws.'),
]

for q, a in qa_pairs:
    add_heading(q, level=3)
    add_paragraph(a)

add_heading('Accounts & Transactions', level=2)

qa_pairs2 = [
    ('Can I have multiple accounts?', 'Yes, unlimited accounts. Each can be a different currency and institution (bank, credit card, investment portfolio).'),
    ('What if I manually enter transactions instead of importing?', 'Perfectly fine. You can log transactions manually. Use for cash expenses, or when bank import isn\'t available.'),
    ('Can I edit or delete transactions?', 'Yes. Click on any transaction → Edit or Delete. History is preserved for audit purposes.'),
    ('What categories should I use?', 'Use what makes sense to you. Common ones: Groceries, Rent, Utilities, Transport, Entertainment, Medical, Education. Create custom categories as needed.'),
    ('How do I handle credit card payments?', 'Credit card payments are "Transfers" not expenses. If you pay your credit card via bank transfer, log it as a transfer between accounts.'),
]

for q, a in qa_pairs2:
    add_heading(q, level=3)
    add_paragraph(a)

add_heading('Multi-Currency & Remittances', level=2)

qa_pairs3 = [
    ('How are exchange rates determined?', 'Rates are fetched from live market sources. You can view the source and see historical rates for any date.'),
    ('Can I lock an exchange rate for a transfer?', 'Yes. When making a transfer, you can enter the rate you actually paid (useful if bank gave you a different rate).'),
    ('How do I track money I send to family?', 'Use the Remittances feature. Log the transfer, upload receipt, and the app tracks efficiency and auto-links to deposits.'),
    ('Can family member receive money to a different account?', 'Yes. Remittance can go to any account (parent, spouse, etc.). Just upload the receipt as proof.'),
]

for q, a in qa_pairs3:
    add_heading(q, level=3)
    add_paragraph(a)

add_heading('Family Features', level=2)

qa_pairs4 = [
    ('Can my spouse see my accounts?', 'Only if you invite them and grant permission. You control access level for each family member.'),
    ('How do we handle shared expenses?', 'Use Shared Budgets. Create a "Household" budget category, and both can log expenses against it.'),
    ('Can I view family members\' full financial details?', 'Only if they grant permission. Privacy controls are per-person and per-section.'),
]

for q, a in qa_pairs4:
    add_heading(q, level=3)
    add_paragraph(a)

add_heading('Loans & Investments', level=2)

qa_pairs5 = [
    ('How do I track an ongoing loan?', 'Go to Loans section, enter loan details (principal, rate, tenure), and the app auto-calculates EMI. Log EMI payments as transactions.'),
    ('What investment types are supported?', 'Stocks, Mutual Funds, FDs, Gold, Crypto, Real Estate, or Other. Manually enter quantities and current values; we calculate gains.'),
    ('Can I import investment statements?', 'Yes. Use the Smart Import feature to upload statements from brokers or asset managers. Estelle AI can extract details.'),
]

for q, a in qa_pairs5:
    add_heading(q, level=3)
    add_paragraph(a)

add_heading('Goals & Budgeting', level=2)

qa_pairs6 = [
    ('How do I set a financial goal?', 'Go to Goals → Add Goal. Enter target amount, deadline, and what you\'re saving for. App calculates required monthly savings.'),
    ('What if I don\'t meet my budget?', 'The app alerts you. You can adjust your budget or track which categories are over and why.'),
    ('Can I have multiple budgets?', 'Yes. Household budget, personal budget, emergency fund, etc. Budgets can overlap categories.'),
]

for q, a in qa_pairs6:
    add_heading(q, level=3)
    add_paragraph(a)

add_heading('AI & Advanced Features', level=2)

qa_pairs7 = [
    ('What can Estelle (AI) do?', 'Ask financial questions, get spending insights, plan scenarios, and extract data from documents (statements, loan papers, receipts).'),
    ('Is my data used to train the AI?', 'No. Your data stays private. Estelle uses OpenAI/Anthropic APIs but doesn\'t use your data for model training.'),
    ('Can I export a tax report?', 'Yes. Settings → Export → Tax Report. Select country and date range. Get itemized income, deductions, and estimated tax liability.'),
]

for q, a in qa_pairs7:
    add_heading(q, level=3)
    add_paragraph(a)

add_heading('Billing & Premium', level=2)

qa_pairs8 = [
    ('How much does premium cost?', 'Premium plans start at $4.99/month or $49.99/year. See pricing page for current rates and features.'),
    ('Can I cancel anytime?', 'Yes, no lock-in. Cancel subscription anytime from Settings → Billing.'),
    ('What\'s included in premium?', 'Family sharing, advanced tax planning, priority support, unlimited imports, and AI advisor features.'),
    ('Do I get a refund if I cancel?', 'Refunds follow your payment method\'s policies. For questions, contact support.'),
]

for q, a in qa_pairs8:
    add_heading(q, level=3)
    add_paragraph(a)

add_heading('Data, Security & Privacy', level=2)

qa_pairs9 = [
    ('How do I download my data?', 'Settings → Data Management → Export. Choose CSV, Excel, or PDF. All your data is yours.'),
    ('Can I delete my account?', 'Yes, Settings → Account → Delete Account. This permanently deletes all data. Action cannot be undone.'),
    ('Who has access to my data?', 'Only you (and family members you invite). Our team doesn\'t access your financial data.'),
    ('Is the app GDPR compliant?', 'Yes, and we comply with privacy laws in all supported countries.'),
]

for q, a in qa_pairs9:
    add_heading(q, level=3)
    add_paragraph(a)

add_paragraph(' ')

# ============= SECTION 11: OFFBOARDING =============
add_heading('11. Offboarding & Data Export', level=1)

add_heading('11.1 Before You Go', level=2)
add_paragraph('If you\'re planning to leave the app, here\'s what to do:')

add_numbered([
    'Export all your data (Settings → Data Management → Export).',
    'Download as Excel or CSV for your records.',
    'Export tax reports for current year.',
    'Save any important documents or attachments.',
    'Share family access if needed before offboarding.',
])

add_heading('11.2 Export Formats Available', level=2)

table = add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Format'
header_cells[1].text = 'Contents'

# Data rows
row_data = [
    ('Excel (.xlsx)', 'All accounts, transactions, balances, remittances, loans, investments'),
    ('CSV (.csv)', 'Data in comma-separated format for import to other tools'),
    ('PDF Report', 'Formatted financial summary with charts and insights'),
]

for i, (fmt, contents) in enumerate(row_data, 1):
    cells = table.rows[i].cells
    cells[0].text = fmt
    cells[1].text = contents

add_heading('11.3 Deleting Your Account', level=2)
add_paragraph('Ready to close your account?')
add_numbered([
    'Go to Settings → Account Management → Delete Account.',
    'Read the warning (this cannot be undone).',
    'Confirm by entering your password.',
    'Account is permanently deleted with all associated data.',
])

add_paragraph('Note: Any shared access with family will also be revoked.')

add_heading('11.4 Switching to Another App', level=2)
add_paragraph('Migrating to a different financial app?')
add_numbered([
    'Export your data from NRI Finance App (CSV or Excel).',
    'Check if the new app supports import in similar format.',
    'Import your data (format may need adjustment).',
    'Verify balances and transactions match.',
    'Then delete your NRI Finance account if no longer needed.',
])

add_heading('11.5 Taking a Break', level=2)
add_paragraph('Don\'t want to delete but taking a break?')
add_bullets([
    'Keep your account inactive (no charge for free tier)',
    'Your data will be preserved indefinitely',
    'Reactivate anytime by logging back in',
])

add_paragraph(' ')

# ============= SECTION 12: SUPPORT & RESOURCES =============
add_heading('12. Contact & Support', level=1)

add_heading('12.1 Get Help', level=2)
add_bullets([
    'In-app Help: Click "?" in the top-right corner',
    'Email: support@nrifinance.app',
    'Chat: Live chat during business hours (9 AM - 6 PM IST)',
    'Community Forum: Share tips and ask questions',
    'Video Tutorials: YouTube channel with feature walkthroughs',
])

add_heading('12.2 Report a Bug', level=2)
add_numbered([
    'Go to Settings → Help → Report Issue.',
    'Describe the problem and steps to reproduce.',
    'Include screenshots if possible.',
    'We\'ll investigate and get back to you within 24 hours.',
])

add_heading('12.3 Suggest a Feature', level=2)
add_paragraph('Have an idea? We love feedback!')
add_numbered([
    'Settings → Help → Feature Request.',
    'Describe what you need and why.',
    'Vote on other requests in the community.',
    'We review top requests quarterly.',
])

add_heading('12.4 Useful Links', level=2)
add_bullets([
    'Website: https://www.nrifinance.app',
    'Blog: https://blog.nrifinance.app',
    'Knowledge Base: https://help.nrifinance.app',
    'Privacy Policy: https://www.nrifinance.app/privacy',
    'Terms of Service: https://www.nrifinance.app/terms',
])

add_heading('12.5 Social Media', level=2)
add_bullets([
    'Twitter: @NRIFinanceApp',
    'Instagram: @nrifinance',
    'LinkedIn: linkedin.com/company/nrifinanceapp',
])

add_paragraph(' ')

# ============= APPENDIX =============
add_heading('Appendix: Workflow Diagrams & Flowcharts', level=1)

add_heading('A. User Onboarding Flowchart', level=2)
add_code('START\n  ↓\n[Sign Up] → Enter Email & Password\n  ↓\n[Email Verification] ← Check Email for Link\n  ↓\n[Login]\n  ↓\n[Setup Wizard]\n  ├─ Select Home Country\n  ├─ Select Working Country\n  ├─ Choose Primary Currency\n  ├─ Create First Account\n  └─ Set Preferences\n  ↓\n[Dashboard Loaded]\n  ↓\nEND (Ready to Use)')

add_heading('B. Adding a Transaction', level=2)
add_code('[Click Add Transaction]\n  ↓\n[Choose Type] → Expense/Income/Transfer\n  ↓\n[Enter Details]\n  ├─ Amount\n  ├─ Category\n  ├─ Date\n  ├─ Account\n  └─ Notes\n  ↓\n[Optional: Attach Receipt]\n  ↓\n[Save]\n  ↓\n[Transaction Added to Ledger]\n  ↓\n[Dashboard Updated]\n  ↓\nEND')

add_heading('C. Remittance Workflow', level=2)
add_code('[User Sends Money Abroad]\n  ↓\n[Log Remittance]\n  ├─ Amount\n  ├─ Exchange Rate\n  ├─ Recipient\n  └─ Upload Receipt\n  ↓\n[Track in Remittances Section]\n  ↓\n[Family Deposits in India]\n  ↓\n[App Auto-Links Transfer]\n  ↓\n[Show Efficiency]\n  ├─ Amount Sent\n  ├─ Amount Received\n  └─ Conversion Loss\n  ↓\nEND')

add_heading('D. Import & Categorize Workflow', level=2)
add_code('[Click Import Bank Statement]\n  ↓\n[Upload File] → PDF/Excel/CSV/Photo\n  ↓\n[System Extracts Transactions]\n  ↓\n[Auto-Categorize]\n  ├─ Match with history\n  ├─ Detect transfers\n  └─ Detect anomalies\n  ↓\n[User Reviews]\n  ├─ Verify categories\n  ├─ Confirm amounts\n  └─ Flag discrepancies\n  ↓\n[Confirm Import]\n  ↓\n[Transactions Added]\n  ↓\nEND')

add_heading('E. Account Reconciliation Flowchart', level=2)
add_code('[Open Account]\n  ↓\n[Click Audit Balance]\n  ↓\n[Enter Bank Statement Balance]\n  ↓\n[System Calculates]\n  ├─ App Balance vs. Bank Balance\n  ├─ Difference\n  └─ Pending Transactions\n  ↓\n[Compare]\n  ├─ Match? → END (Reconciled)\n  └─ Mismatch? ↓\n      ↓\n      [Find Missing Transactions]\n      ├─ Check uncleared items\n      ├─ Check categorization\n      └─ Check date\n      ↓\n      [Add/Edit Transactions]\n      ↓\n      [Re-reconcile]\n      ↓\n      END')

add_heading('F. Offboarding & Data Export', level=2)
add_code('[User Decides to Leave]\n  ↓\n[Settings → Data Management]\n  ↓\n[Choose Export Format]\n  ├─ Excel\n  ├─ CSV\n  └─ PDF Report\n  ↓\n[Select Date Range]\n  ↓\n[Download File]\n  ↓\n[Settings → Delete Account]\n  ↓\n[Confirm Password]\n  ↓\n[Account & All Data Deleted]\n  ↓\nEND')

add_paragraph(' ')
add_paragraph('=' * 80)
add_paragraph('End of User Manual')
add_paragraph('For the latest updates and new features, visit: https://www.nrifinance.app')
add_paragraph('Last Updated: June 2026')
add_paragraph('Version: 1.0')

# Save the document
output_path = 'NRI_Finance_App_User_Manual.docx'
doc.save(output_path)
print(f'✓ User Manual created: {output_path}')
