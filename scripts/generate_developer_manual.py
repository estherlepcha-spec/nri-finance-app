from docx import Document
from docx.shared import Pt


doc = Document()
doc.core_properties.title = 'NRI Finance App Developer Manual'
doc.core_properties.subject = 'Developer manual, product specification, workflows, and troubleshooting guide.'
doc.core_properties.author = 'NRI Finance App Team'

def add_heading(text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(text, style=None):
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    return p


def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_numbered(items):
    for item in items:
        doc.add_paragraph(item, style='List Number')


def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    font = run.font
    font.name = 'Courier New'
    font.size = Pt(10)
    p.paragraph_format.left_indent = Pt(12)
    return p

# Title
add_heading('NRI Finance App Developer Manual', level=0)
add_paragraph('Complete product specification, feature overview, implementation logic, workflows, troubleshooting, and go-to-market considerations.')

# Section 1: Product overview
add_heading('1. Product Overview', level=1)
add_paragraph('The NRI Finance App is a personal finance manager designed for Non-Resident Indians and expatriates. It supports multi-currency finance tracking, remittances, loans, investments, budgets, and AI-assisted financial insights.')
add_bullets([
    'Target users: NRIs, expats, and globally-mobile professionals.',
    'Platforms: Web application built with React + Vite.',
    'Persistence: LocalStorage plus Supabase cloud sync and optional real-time sync.',
    'AI features: Claude-powered document extraction and advisor workflows.',
    'Core differentiators: multi-currency remittance tracking, loan/EMI intelligence, and family finance support.',
])

# Section 2: Features and product specification
add_heading('2. Product Specification and Features', level=1)
add_heading('2.1 Core Features', level=2)
add_bullets([
    'Dashboard overview with net worth, monthly income/expense summary, and country breakdown.',
    'Account management for home and working country accounts with currency-aware balances.',
    'Transaction ledger with smart categorization and bank statement import support.',
    'Remittance tracking with transfer receipts, exchange rate conversion, and auto-linking.',
    'Loan and EMI tracking with payoff estimates and loan matching from transaction history.',
    'Investment portfolio tracking for mutual funds, stocks, and foreign assets.',
    'Bill reminder and recurring payment tracking.',
    'Goals and budget management for savings planning.',
    'AI assistant (Estelle) for financial advice and document parsing.',
])

add_heading('2.2 Additional Capabilities', level=2)
add_bullets([
    'Multi-currency support with INR as the primary home currency and foreign working currencies.',
    'Balance audit modal to reconcile account balance, transaction history, and remittance receipts.',
    'Real-time or near-real-time sync via Supabase and optional Vite dev sync endpoint.',
    'Custom rule and import cleanup support for duplicate detection and intelligent transaction classification.',
    'Support for family members, shared budgets, and auxiliary accounts.',
])

add_heading('2.3 Key Product Metrics', level=2)
add_bullets([
    'Total assets and liabilities by currency.',
    'Monthly income, expenses, savings, and remittances.',
    'Credit card debt and loan outstanding balances.',
    'Remittance conversion efficiency and missing transaction alerts.',
])

# Section 3: Architecture and modules
add_heading('3. Architecture and Technical Stack', level=1)
add_paragraph('The app uses a modular front-end architecture with React, Vite, and Supabase integration. Key folders and components are listed below.')
add_heading('3.1 Tech Stack', level=2)
add_bullets([
    'Frontend: React 18 + Vite.',
    'Styling: Inline CSS in JavaScript with theme constants and dedicated style files.',
    'Persistence: Supabase PostgreSQL for cloud storage and authentication.',
    'AI: Anthropic Claude via Supabase Edge Functions proxy.',
    'Build tooling: Vite, ESLint, Playwright, Puppeteer.',
])

add_heading('3.2 Core Folder Structure', level=2)
add_code('src/\n  App.jsx\n  main.jsx\n  auth.js\n  sync.js\n  supabase.js\n  services/anthropic.js\n  utils/calculations.js\n  utils/constants.js\n  utils/formatting.jsx\n  components/shared/\n  components/SetupWizard/\n  components/Family/\n')

add_heading('3.3 Data Model and State', level=2)
add_paragraph('App state is primarily managed inside App.jsx using React useState hooks. Data entities include accounts, transactions, remittances, bills, investments, loans, goals, allocations, family, and layouts.')
add_bullets([
    'accounts: list of bank, savings, credit card, and loan accounts with currency and balance metadata.',
    'transactions: ledger entries including type, category, amount, denomination, and linked account.',
    'remittances: money transfers with provider details, received values, and FX rates.',
    'loans: outstanding balances, EMI, interest rate, tenure, and currency.',
    'goals: savings and financial goals with progress tracking.',
])

# Section 4: Function mechanisms and logic
add_heading('4. Functional Logic and Mechanics', level=1)
add_heading('4.1 Balance and Transaction Calculations', level=2)
add_paragraph('The src/utils/calculations.js module contains reusable pure functions for financial computations. These functions ensure balance accuracy across account types and transaction categories.')
add_numbered([
    'calcTxDelta(t, isCC): computes how a transaction affects balance.',
    'getAccountBalanceAtDate(accs, txs, accountId, date): returns cumulative balance up to a date.',
    'getOpeningBalance(accs, txs, accountId, month): returns account balance before month start.',
    'getClosingBalance(accs, txs, accountId, month): returns account balance at month end.',
    'recomputeAllBalances(accs, txs): recomputes live balances from setup balances and transactions.',
    'convertAmountToINR(amount, currency, rates, fallbackExchangeRate): converts any amount to INR.',
    'calculateBalanceAudit(account, txs, remittances, homeCurrency): builds the audit reconciliation view for an account.',
])

add_paragraph('This function set distinguishes normal accounts from credit cards, where income and expense semantics are reversed for balance calculation.')
add_code('def calcTxDelta(t, isCC):\n    amount = abs(t.amount or 0)\n    if isCC:\n        return -amount if t.type == "income" else amount\n    return amount if t.type == "income" else -amount\n')

add_heading('4.2 Account Audit and Reconciliation', level=2)
add_paragraph('The Audit modal in App.jsx uses shared helper logic to display a breakdown of the account balance, transaction deltas, and remittance matching. It detects unlinked remittances and estimates expected balance.')
add_paragraph('Workflow:')
add_code('1. Load account and transactions for selected account.\n2. Determine if account is a credit card.\n3. Build transaction deltas with calcTxDelta.\n4. Sum increases and decreases.\n5. Recompute app balance from setup balance + deltas.\n6. Filter remittances for the account currency.\n7. Match remittances to income transactions using tolerance rules.\n8. Compute unlinked remittance total and expected balance.\n')

add_heading('4.3 AI Integration', level=2)
add_paragraph('The AI workflow is implemented in src/services/anthropic.js. The app calls Supabase Edge Functions to keep the Claude API key secure on the server side.')
add_bullets([
    'anthropicMessages(body): low-level proxy call to the Supabase Edge Function.',
    'callClaude(messages, options): sends a standard Anthropic message payload.',
    'extractWithPrompt(prompt, content, maxTokens): helper for text prompt extraction.',
    'extractFromFile(prompt, fileContent, fileType, maxTokens): handles PDF, image, and text payloads.',
])
add_paragraph('The AI proxy uses the authenticated user session token so the browser never holds the Claude secret key.')

add_heading('4.4 Persistence and Sync', level=2)
add_paragraph('Persistence is handled by src/supabase.js and optional local sync layers in src/sync.js. The app persists selected keys to Supabase and can subscribe to realtime updates.')
add_numbered([
    'loadFromSupabase(): load user-specific data rows from the nri_finance_data table.',
    'saveToSupabase(key, value): upsert data row for the signed-in user.',
    'subscribeToChanges(onUpdate): listen to Postgres changes for the user.',
    'sync.js init(onStatus, onRemote): connect to the /api/sync endpoint and push local changes.',
])

add_heading('4.5 Local Storage and Remote Sync', level=2)
add_paragraph('The app caches state in localStorage and sends remote sync writes when available. The sync layer avoids empty server-state overwrites when the user is signed in.')
add_code('function applyFiltered(data, isInit = false):\n    for each sync key:\n        if isInit and server value is empty: keep local value if present\n        otherwise update local state\n')

# Section 5: Workflows and diagrams
add_heading('5. Workflows and Diagrams', level=1)
add_heading('5.1 Startup Data Load Workflow', level=2)
add_code('App Start -> read localStorage -> set default state -> if user signed in then\n  loadFromSupabase() -> merge with local state -> render UI\n')
add_paragraph('Explanation: the app always initializes from local cache first and then hydrates from Supabase when authentication is available.')

add_heading('5.2 Transaction Add / Recompute Workflow', level=2)
add_code('User adds transaction -> create tx object -> update transactions state ->\n  setAccounts(recomputeAllBalances(accounts, transactions)) -> persist() -> remote sync or Supabase save\n')
add_paragraph('This ensures account balances remain consistent whenever transaction data changes.')

add_heading('5.3 Balance Audit Workflow', level=2)
add_code('Open Audit modal -> calculateBalanceAudit(account, transactions, remittances) -> show app balance vs expected balance -> optionally sync unlinked remittances -> update transactions and balances\n')
add_paragraph('The audit workflow helps the user understand why an account balance is out of sync with transfers and remittance receipts.')

add_heading('5.4 AI Document Extraction Workflow', level=2)
add_code('User uploads a document or text -> prepare prompt -> call extractFromFile() or extractWithPrompt() -> proxy via Supabase function -> parse response -> map extracted fields to transactions / remittances / investments\n')

add_heading('5.5 Real-Time Sync Workflow', level=2)
add_code('Sync init -> verify /api/sync endpoint -> open EventSource -> receive init/update events -> applyFiltered() -> push local updates through POST -> handle reconnects\n')

# Section 6: Troubleshooting
add_heading('6. Troubleshooting Guide', level=1)
add_heading('6.1 Build and Dependency Issues', level=2)
add_bullets([
    'If npm run build fails, ensure dependencies are installed with npm install and that vite is the correct version.',
    'Check package.json for React, Vite, and supabase versions; mismatched peer dependencies may break build.',
    'Large chunks warning is informational; use code-splitting or dynamic imports to reduce bundle size if needed.',
    'If React or jsx syntax errors appear, confirm @vitejs/plugin-react is enabled in vite.config.js.',
])

add_heading('6.2 Supabase and Authentication Issues', level=2)
add_bullets([
    'If Supabase loads fail, verify VITE_SUPABASE_URL and VITE_SUPABASE_ANON_KEY in .env or environment config.',
    'Ensure the Supabase table nri_finance_data exists and has user_id, key, value, and updated_at fields.',
    'If auth session is null, calls to supabase.auth.getSession() return no user; users must sign in before cloud persistence works.',
    'Inspect network requests for functions/v1/anthropic and api/sync to identify authorization or endpoint issues.',
])

add_heading('6.3 Data Sync and State Issues', level=2)
add_bullets([
    'If local state does not match Supabase, verify that keys in SYNC_KEYS are consistent across src/supabase.js and src/sync.js.',
    'Empty server state should not overwrite local cache due to the init filtering logic; if it does, check applyFiltered() implementation.',
    'For stale UI updates, confirm setAccounts(recomputeAllBalances(...)) is called after transactions change.',
    'If duplicate transactions appear, inspect import logic and smart rules for duplicate detection. Remove duplicate keys manually if needed.',
])

add_heading('6.4 Audit Modal Troubleshooting', level=2)
add_bullets([
    'If the audit balance is incorrect for credit cards, verify isCC is set and calcTxDelta uses reversed semantics for income vs expense.',
    'If remittances are not linked, examine the tolerance rule in calculateBalanceAudit and the toCurrency matching logic.',
    'If expected balance does not update after syncing, confirm setTransactions(updated) and setAccounts(prev => recomputeAllBalances(prev, updated)) execute correctly.',
])

add_heading('6.5 AI and Claude Issues', level=2)
add_bullets([
    'If AI extraction fails, ensure the Supabase Edge Function is deployed and the VITE_SUPABASE_URL + VITE_SUPABASE_ANON_KEY are correct.',
    'Check that the user session exists before making anthropicMessages() calls; unauthorized requests will fail.',
    'If the proxy returns JSON parse errors, inspect the edge function response format and message payload structure.',
])

# Section 7: Selling the App
add_heading('7. Preparing the App for Sale', level=1)
add_heading('7.1 Product Readiness', level=2)
add_bullets([
    'Complete feature polish and bug fixes, especially around data accuracy, remittances, and multi-currency conversions.',
    'Add onboarding and help documentation for target users.',
    'Ensure the app is responsive, accessible, and user-friendly across devices.',
])

add_heading('7.2 Legal and Compliance', level=2)
add_bullets([
    'Prepare Terms of Service and Privacy Policy documents.',
    'Ensure GDPR and data privacy compliance for user financial data.',
    'If you plan to process payments or subscriptions, comply with local financial regulations.',
    'Consider consulting a lawyer for global user licensing and data storage laws.',
])

add_heading('7.3 Deployment and Hosting', level=2)
add_bullets([
    'Use Vercel or another static host for the frontend.',
    'Host Supabase on a production project with a secure database and authentication.',
    'Set up environment variables securely and use separate prod/dev Supabase projects.',
    'Configure monitoring, logging, and backup policies for the database.',
])

add_heading('7.4 Monetization Options', level=2)
add_paragraph('Common monetization strategies include:')
add_bullets([
    'Subscription plans for premium features, AI advisor, sync, and advanced analytics.',
    'One-time purchase for the app or a paid SaaS product.',
    'Freemium model with core features free and paid upgrades.',
    'Partner programs with financial institutions or remittance firms.',
])

add_heading('7.5 Sales and Go-To-Market Needs', level=2)
add_bullets([
    'Branding, product website, and marketing collateral.',
    'Demo videos, tutorials, and customer support flows.',
    'A payment gateway integration for subscriptions or one-time sales.',
    'Legal entity setup, accounting, and tax registration.',
])

add_heading('7.6 Scaling and Support', level=2)
add_bullets([
    'Plan for customer support and bug triage workflows.',
    'Use crash reporting and analytics to identify adoption and retention issues.',
    'Prepare documentation for onboarding new developers or partners.',
])

# Final section
add_heading('8. Appendix: Key Implementation Files', level=1)
add_paragraph('This appendix lists the most important files and their responsibilities.')
add_numbered([
    'src/App.jsx – root application state, page layout, feature integrations, audit modal, reconcile flow.',
    'src/utils/calculations.js – shared financial helper functions and audit calculation logic.',
    'src/services/anthropic.js – AI proxy service for Claude API requests.',
    'src/supabase.js – Supabase client, auth helper, persistence, and real-time change subscription.',
    'src/sync.js – local real-time sync layer for development sync endpoint and push/pull updates.',
    'src/components/shared/index.jsx – reusable UI components such as buttons, modals, cards, and inputs.',
])

path = 'NRI_Finance_App_Developer_Manual.docx'
doc.save(path)
print(path)
